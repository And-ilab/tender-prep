#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Извлечение текста из локального корпуса (после corpus-pull и т.п.):
  .docx  -> python-docx; при невалидном OOXML или ошибке — LibreOffice (если есть)
  .pdf   -> pymupdf (опционально)
  .doc   -> LibreOffice headless (если soffice в PATH или --libreoffice)
  .txt/.md/.csv/.log -> копия в выходное дерево как .txt

Исходная папка --input не изменяется. В --output зеркало относительных путей + .txt
"""

from __future__ import annotations

import argparse
import glob
import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

SKIP_DIR_NAMES = {".git", "node_modules", "__pycache__", ".venv", ".lena-rag"}
TEXT_COPY_EXT = {".txt", ".md", ".markdown", ".csv", ".log"}

SCRIPT_VERSION = 3


@dataclass
class Report:
    copied_text: int = 0
    docx_ok: int = 0
    docx_via_libreoffice: int = 0
    docx_err: list[dict[str, str]] = field(default_factory=list)
    pdf_ok: int = 0
    pdf_skip: int = 0
    pdf_err: list[dict[str, str]] = field(default_factory=list)
    doc_ok: int = 0
    doc_skip: int = 0
    doc_err: list[dict[str, str]] = field(default_factory=list)
    skipped_other: int = 0
    other_err: list[dict[str, str]] = field(default_factory=list)


def find_soffice(explicit: str | None) -> str | None:
    """Путь к soffice.exe: аргумент CLI, переменные окружения, PATH, стандартные каталоги Windows."""
    if explicit:
        p = Path(explicit)
        return str(p) if p.is_file() else None

    for env_key in ("LENA_LIBREOFFICE_SOFFICE", "LIBREOFFICE_SOFFICE"):
        raw = (os.environ.get(env_key) or "").strip().strip('"')
        if raw:
            ep = Path(raw)
            if ep.is_file() and ep.name.lower().startswith("soffice"):
                return str(ep)
            if ep.is_dir():
                cand = ep / "soffice.exe"
                if cand.is_file():
                    return str(cand)

    for name in ("soffice", "soffice.exe"):
        found = shutil.which(name)
        if found:
            return found

    for pattern in (
        r"C:\Program Files\LibreOffice*\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice*\program\soffice.exe",
    ):
        matches = sorted(glob.glob(pattern), reverse=True)
        for m in matches:
            if Path(m).is_file():
                return m

    win = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    if win.is_file():
        return str(win)
    win32 = Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe")
    if win32.is_file():
        return str(win32)
    return None


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        rows_out: list[str] = []
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            rows_out.append("\t".join(cells))
        if rows_out:
            parts.append("\n".join(rows_out))
    return "\n\n".join(parts).strip()


def is_ooxml_word_package(path: Path) -> bool:
    """True только если внутри ZIP есть типичная разметка WordprocessingML."""
    try:
        with zipfile.ZipFile(path, "r") as zf:
            names = [n.lower() for n in zf.namelist()]
    except zipfile.BadZipFile:
        return False
    if "[content_types].xml" in names:
        return True
    if "word/document.xml" in names:
        return True
    return False


def try_extract_docx_python(path: Path) -> tuple[str, str | None]:
    """
    Возвращает (код, текст):
    - ("ok", str) — успех python-docx, строка может быть пустой
    - ("not_ooxml", None) — не ZIP / не пакет Word OOXML
    - ("broken_ooxml", None) — пакет похож на OOXML, но Document() не читает
    """
    if not path.is_file():
        return "not_ooxml", None
    if path.stat().st_size == 0:
        return "ok", ""
    if not zipfile.is_zipfile(path):
        return "not_ooxml", None
    if not is_ooxml_word_package(path):
        return "not_ooxml", None
    try:
        return "ok", extract_docx(path)
    except Exception:
        return "broken_ooxml", None


def extract_pdf(path: Path) -> str:
    import fitz  # pymupdf

    text_parts: list[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            t = page.get_text("text") or ""
            t = t.strip()
            if t:
                text_parts.append(t)
    return "\n\n".join(text_parts).strip()


def convert_doc_with_libreoffice(soffice: str, src: Path, dest_txt: Path) -> None:
    dest_txt.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="lo_conv_") as tmp:
        tmp_path = Path(tmp)
        cp = subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "txt",
                "--outdir",
                str(tmp_path),
                str(src.resolve()),
            ],
            check=False,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
            text=True,
            timeout=600,
        )
        if cp.returncode != 0:
            err = (cp.stderr or "").strip()[:1200]
            raise RuntimeError(f"LibreOffice exit {cp.returncode}. stderr: {err or '(пусто)'}")
        produced = tmp_path / (src.stem + ".txt")
        if not produced.is_file():
            raise RuntimeError(f"LibreOffice не создал {produced.name} в {tmp_path}")
        shutil.copy2(produced, dest_txt)


def safe_rel(path: Path, root: Path) -> str:
    return path.relative_to(root).as_posix()


def mirror_out_path(in_root: Path, out_root: Path, src_file: Path) -> Path:
    rel = src_file.relative_to(in_root)
    stem = rel.stem
    parent_parts = rel.parts[:-1]
    return out_root.joinpath(*parent_parts, f"{stem}.txt")


def walk_files(in_root: Path) -> list[Path]:
    files: list[Path] = []
    for p in in_root.rglob("*"):
        if not p.is_file():
            continue
        if any(part in SKIP_DIR_NAMES for part in p.parts):
            continue
        files.append(p)
    files.sort()
    return files


def process_one(
    src: Path,
    in_root: Path,
    out_root: Path,
    args: argparse.Namespace,
    has_fitz: bool,
    soffice: str | None,
    rep: Report,
) -> None:
    rel = safe_rel(src, in_root)
    ext = src.suffix.lower()

    if ext in TEXT_COPY_EXT:
        dest = mirror_out_path(in_root, out_root, src)
        if dest.exists() and not args.force:
            rep.copied_text += 1
            return
        dest.parent.mkdir(parents=True, exist_ok=True)
        text = src.read_text(encoding="utf-8", errors="replace")
        dest.write_text(text, encoding="utf-8", newline="\n")
        rep.copied_text += 1
        return

    if ext == ".docx":
        dest = mirror_out_path(in_root, out_root, src)
        if dest.exists() and not args.force:
            rep.docx_ok += 1
            return
        dest.parent.mkdir(parents=True, exist_ok=True)
        mode, text = try_extract_docx_python(src)
        if mode == "ok":
            dest.write_text(text if text else "(пустой документ)\n", encoding="utf-8", newline="\n")
            rep.docx_ok += 1
            return
        if soffice:
            convert_doc_with_libreoffice(soffice, src, dest)
            rep.docx_via_libreoffice += 1
            rep.docx_ok += 1
            return
        hint = (
            "Файл не является валидным OOXML (.docx как ZIP) или повреждён; "
            "python-docx не открыл. Установите LibreOffice и добавьте soffice в PATH "
            "или укажите --libreoffice для конвертации."
        )
        if mode == "not_ooxml":
            hint = (
                "Не найден ZIP-контейнер Office Open XML (часто: нулевой размер, "
                "обрыв выгрузки с Drive, или под видом .docx лежит HTML/ярлык). " + hint
            )
        raise RuntimeError(hint)

    if ext == ".pdf":
        if not has_fitz:
            rep.pdf_skip += 1
            return
        dest = mirror_out_path(in_root, out_root, src)
        if dest.exists() and not args.force:
            rep.pdf_ok += 1
            return
        dest.parent.mkdir(parents=True, exist_ok=True)
        body = extract_pdf(src)
        dest.write_text(body if body else "(PDF без извлекаемого текста)\n", encoding="utf-8", newline="\n")
        rep.pdf_ok += 1
        return

    if ext == ".doc":
        if not soffice:
            rep.doc_skip += 1
            return
        dest = mirror_out_path(in_root, out_root, src)
        if dest.exists() and not args.force:
            rep.doc_ok += 1
            return
        convert_doc_with_libreoffice(soffice, src, dest)
        rep.doc_ok += 1
        return

    rep.skipped_other += 1


def main() -> int:
    ap = argparse.ArgumentParser(description="Корпус -> извлечённый текст (UTF-8), без удаления оригиналов")
    ap.add_argument("--input", "-i", type=Path, required=True, help="Корень выгрузки corpus-pull")
    ap.add_argument("--output", "-o", type=Path, required=True, help="Корень для зеркала .txt")
    ap.add_argument("--force", action="store_true", help="Перезаписать существующие .txt")
    ap.add_argument("--dry-run", action="store_true", help="Только статистика по файлам")
    ap.add_argument("--no-pdf", action="store_true", help="Не обрабатывать PDF")
    ap.add_argument("--libreoffice", type=str, default=None, help="Путь к soffice.exe: .doc и запасной путь для «битых» .docx")
    ap.add_argument(
        "--lenient-exit",
        action="store_true",
        help="Всегда код выхода 0 после записи report.json (ошибки только в JSON отчёта; для CI/следующего шага rag index)",
    )
    ap.add_argument("--version", action="store_true", help="Печать версии скрипта и выход")
    args = ap.parse_args()

    if args.version:
        print(SCRIPT_VERSION)
        return 0

    in_root: Path = args.input.resolve()
    out_root: Path = args.output.resolve()
    if not in_root.is_dir():
        print(f"Нет папки input: {in_root}", file=sys.stderr)
        return 2

    has_fitz = False
    if not args.no_pdf:
        try:
            import fitz  # noqa: F401

            has_fitz = True
        except ImportError:
            pass

    try:
        import docx  # noqa: F401
    except ImportError:
        print("Нужен пакет python-docx: pip install -r requirements.txt", file=sys.stderr)
        return 2

    soffice = find_soffice(args.libreoffice)
    print(
        f"[corpus_extract_text] script_version={SCRIPT_VERSION} LibreOffice: {soffice or 'НЕ НАЙДЕН (задайте LENA_LIBREOFFICE_SOFFICE или --libreoffice)'}",
        file=sys.stderr,
    )
    files = walk_files(in_root)
    by_ext: dict[str, int] = {}
    for f in files:
        by_ext[f.suffix.lower()] = by_ext.get(f.suffix.lower(), 0) + 1

    if args.dry_run:
        print(
            json.dumps(
                {
                    "script_version": SCRIPT_VERSION,
                    "input": str(in_root),
                    "file_count": len(files),
                    "by_extension": by_ext,
                    "libreoffice": soffice,
                    "pymupdf": has_fitz,
                },
                ensure_ascii=False,
                indent=2,
            )
        )
        return 0

    out_root.mkdir(parents=True, exist_ok=True)
    rep = Report()

    for src in files:
        rel = safe_rel(src, in_root)
        try:
            process_one(src, in_root, out_root, args, has_fitz, soffice, rep)
        except Exception as e:
            ext = src.suffix.lower()
            err = {"path": rel, "error": str(e)[:800]}
            if ext == ".docx":
                rep.docx_err.append(err)
            elif ext == ".pdf":
                rep.pdf_err.append(err)
            elif ext == ".doc":
                rep.doc_err.append(err)
            else:
                rep.other_err.append(err)

    report_path = out_root / "report.json"
    payload: dict[str, Any] = {
        "script_version": SCRIPT_VERSION,
        "input": str(in_root),
        "output": str(out_root),
        "libreoffice": soffice,
        "pymupdf": has_fitz,
        "copied_text_files": rep.copied_text,
        "docx_extracted": rep.docx_ok,
        "docx_via_libreoffice_fallback": rep.docx_via_libreoffice,
        "docx_errors": rep.docx_err,
        "pdf_extracted": rep.pdf_ok,
        "pdf_skipped_no_pymupdf": rep.pdf_skip,
        "pdf_errors": rep.pdf_err,
        "doc_libreoffice_ok": rep.doc_ok,
        "doc_skipped_no_libreoffice": rep.doc_skip,
        "doc_errors": rep.doc_err,
        "skipped_unhandled_extension": rep.skipped_other,
        "other_errors": rep.other_err,
    }
    report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    print(f"\nОтчёт записан: {report_path}", file=sys.stderr)
    bad = rep.docx_err or rep.pdf_err or rep.doc_err or rep.other_err
    if args.lenient_exit:
        return 0
    return 1 if bad else 0


if __name__ == "__main__":
    sys.exit(main())
